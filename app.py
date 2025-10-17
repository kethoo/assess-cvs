import streamlit as st
from cv_assessment import CVAssessmentSystem
import tempfile
import os
import re
import pandas as pd
import json
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ------------------- STREAMLIT CONFIG -------------------

st.set_page_config(page_title="Deep CV Assessment System", layout="wide")
st.title("📄 Deep CV Assessment System")

# ------------------- SESSION STATE INITIALIZATION -------------------

if 'expert_section_text' not in st.session_state:
    st.session_state.expert_section_text = ""

if 'custom_criteria' not in st.session_state:
    st.session_state.custom_criteria = None

if 'criteria_generated' not in st.session_state:
    st.session_state.criteria_generated = False

if 'editable_df' not in st.session_state:
    st.session_state.editable_df = None

# ------------------- API KEY & EVALUATOR INFO -------------------

import os
import streamlit as st
from dotenv import load_dotenv

# Try to load API key securely
load_dotenv()  # Loads from .env if running locally
api_key = os.getenv("OPENAI_API_KEY", None)

# Fallback: use Streamlit Secrets if available (for Streamlit Cloud)
if not api_key and "OPENAI_API_KEY" in st.secrets:
    api_key = st.secrets["OPENAI_API_KEY"]

# Safety check
if not api_key:
    st.error("❌ No API key found. Please set OPENAI_API_KEY in Streamlit Secrets or a .env file.")
else:
    st.success("🔐 API key loaded securely.")

evaluator_name = st.text_input("👤 Enter Evaluator Name (for report cover page)", placeholder="e.g. John Smith")

# ------------------- UPLOAD TENDER -------------------

req_file = st.file_uploader("📄 Upload Tender / Job Description", type=["pdf", "docx", "doc"])
tender_text = ""
tender_path = None

if req_file:
    suffix = os.path.splitext(req_file.name)[1] or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(req_file.read())
        tender_path = tmp.name

    st.success(f"✅ Tender uploaded: {req_file.name}")

    system_temp = CVAssessmentSystem(api_key=api_key or None)
    tender_text = system_temp.load_job_requirements(tender_path)
    st.info("📘 Tender text loaded successfully.")

# ------------------- ROLE FOCUS SELECTION -------------------

st.markdown("### 🎯 Select Evaluation Focus")
role_focus = st.radio(
    "Choose the type of assessment focus:",
    ["Specific Role (80/20 weighting)", "General Role (100% general weighting)"],
    index=0,
)

# ------------------- SUPPORT FUNCTIONS -------------------

def extract_expert_section_llm(full_text: str, expert_name: str, api_key: str) -> str:
    """Extract specific expert section using LLM."""
    if not full_text or not expert_name or not api_key:
        return ""
    try:
        import openai
        client = openai.OpenAI(api_key=api_key)
        prompt = f"""Extract all text describing the role, qualifications, and responsibilities of "{expert_name}".
Return full sections with all details, preserving structure.
Separate multiple sections with "---SECTION BREAK---".
Document:
{full_text[:30000]}"""
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": "Extract exactly what's requested, no commentary."},
                      {"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=4000,
        )
        text = resp.choices[0].message.content.strip()
        return text.replace("---SECTION BREAK---", "\n\n" + "-"*60 + "\n\n") if text else ""
    except Exception as e:
        st.error(f"LLM extraction error: {e}")
        return ""


def generate_criteria_and_weights(section_text: str, general_context: str, api_key: str, total_weight: int = 80) -> dict:
    """Generate custom criteria and weights."""
    if not section_text or not api_key:
        return None
    try:
        import openai
        client = openai.OpenAI(api_key=api_key)
        prompt = f"""Generate 6–10 key evaluation criteria (name, weight, rationale) totaling {total_weight}%.
Return only JSON like:
{{"criteria":[{{"name":"...","weight":25,"rationale":"..."}}]}}
TEXT:
{section_text[:8000]}
GENERAL CONTEXT:
{general_context[:2000]}"""
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": "Return only valid JSON."},
                      {"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=2000,
        )
        text = re.sub(r"```(json)?|```", "", resp.choices[0].message.content.strip())
        match = re.search(r"(\{.*\})", text, re.DOTALL)
        data = json.loads(match.group(1)) if match else {}
        total = sum(c['weight'] for c in data.get("criteria", []))
        if abs(total - total_weight) > 2:
            factor = total_weight / total
            for c in data["criteria"]:
                c["weight"] = round(c["weight"] * factor, 1)
        return data
    except Exception as e:
        st.error(f"Error generating criteria: {e}")
        return None

# ------------------- SPECIFIC ROLE MODE -------------------

expert_name = ""

if role_focus == "Specific Role (80/20 weighting)":
    st.markdown("### 🧩 Enter the Expert Role Title")
    expert_name = st.text_input("Example: Key Expert 1, Team Leader", placeholder="Type role title")

    st.info("💡 Use AI to extract the section for this expert.")
    if st.button("🔍 Extract Expert Section (AI)", disabled=not (req_file and api_key and expert_name)):
        with st.spinner("Extracting..."):
            extracted = extract_expert_section_llm(tender_text, expert_name, api_key)
            st.session_state.expert_section_text = extracted or ""
            st.success("✅ Extracted successfully!")

    edited_section = st.text_area("📝 Expert Section Content (editable)", value=st.session_state.expert_section_text, height=400)
    st.session_state.expert_section_text = edited_section

    if st.session_state.expert_section_text and api_key:
        st.markdown("---")
        st.markdown("### 🎯 Generate or Edit Criteria (Specific Role)")
        if st.button("🧠 Generate Criteria (80/20 weighting)"):
            with st.spinner("Generating criteria..."):
                data = generate_criteria_and_weights(st.session_state.expert_section_text, tender_text, api_key, 80)
                if data:
                    st.session_state.custom_criteria = data
                    st.session_state.criteria_generated = True
                    st.session_state.editable_df = pd.DataFrame(data["criteria"])
                    st.success("✅ Criteria generated!")

# ------------------- GENERAL ROLE MODE -------------------

if role_focus == "General Role (100% general weighting)" and api_key and req_file:
    st.markdown("### 🧠 Generate or Edit Criteria (General Role)")
    if st.button("🧠 Generate Criteria (100% weighting)"):
        with st.spinner("Generating general criteria..."):
            data = generate_criteria_and_weights(tender_text, tender_text, api_key, 100)
            if data:
                st.session_state.custom_criteria = data
                st.session_state.criteria_generated = True
                st.session_state.editable_df = pd.DataFrame(data["criteria"])
                st.success("✅ General criteria generated!")

# ------------------- EDIT & SAVE CRITERIA -------------------

if st.session_state.criteria_generated and st.session_state.custom_criteria:
    st.markdown("---")
    st.markdown("### ✏️ Edit Your Criteria Table")
    st.info("You can add, edit, or remove rows. Click 💾 Save when finished.")
    edited_df = st.data_editor(
        st.session_state.editable_df,
        use_container_width=True,
        num_rows="dynamic",
        key="criteria_editor"
    )

    if st.button("💾 Save Final Criteria"):
        st.session_state.editable_df = edited_df
        st.session_state.custom_criteria = {"criteria": edited_df.to_dict(orient="records")}
        st.success("✅ Final criteria saved and ready for assessment!")

# ------------------- UPLOAD CVS -------------------

cv_files = st.file_uploader("👤 Upload Candidate CVs", type=["pdf", "docx", "doc"], accept_multiple_files=True)

# ------------------- RUN ASSESSMENT -------------------

st.markdown("---")
st.markdown("### 🚀 Step 4: Run Assessment")

results = []
ranked = []

if st.button("🚀 Run Assessment") and req_file and cv_files and (api_key or os.getenv("OPENAI_API_KEY")):
    with tempfile.TemporaryDirectory() as tmpdir:
        cv_folder = os.path.join(tmpdir, "cvs")
        os.makedirs(cv_folder, exist_ok=True)
        for file in cv_files:
            path = os.path.join(cv_folder, file.name)
            with open(path, "wb") as f:
                f.write(file.read())

        system = CVAssessmentSystem(api_key=api_key or None)

        if role_focus == "Specific Role (80/20 weighting)":
            combined = (
                f"--- GENERAL TENDER CONTEXT (20%) ---\n\n{tender_text[:5000]}\n\n"
                f"--- SPECIFIC EXPERT REQUIREMENTS (80%) ---\n\n{st.session_state.expert_section_text}"
            )
        else:
            combined = f"--- GENERAL ROLE (100%) ---\n\n{tender_text}"

        system.job_requirements = combined
        criteria = st.session_state.custom_criteria if st.session_state.criteria_generated else None
        results = system.process_cv_folder(cv_folder, mode="critical", custom_criteria=criteria)

        st.success(f"✅ Processed {len(results)} candidate(s)")
        ranked = sorted(results, key=lambda x: x.get("final_score", 0), reverse=True)
        st.table([{"Rank": i + 1, "Candidate": r["candidate_name"], "Final Score": f"{r['final_score']:.2f}"} for i, r in enumerate(ranked)])
        for r in ranked:
            with st.expander(f"{r['candidate_name']} — Critical Evaluation"):
                st.markdown(r["report"])

# ------------------- FINAL COMPARISON SUMMARY -------------------

if results:
    st.markdown("---")
    st.markdown("## 🧾 Final Candidate Comparison Summary")

    summary_rows = []
    for i, r in enumerate(ranked):
        report_text = r["report"]
        strengths_match = re.search(r"Strengths(.*?)Weaknesses", report_text, re.DOTALL | re.IGNORECASE)
        weaknesses_match = re.search(r"Weaknesses(.*?)(?:Tailoring|$)", report_text, re.DOTALL | re.IGNORECASE)
        strengths = strengths_match.group(1).strip().replace("\n", " ")[:200] if strengths_match else "-"
        weaknesses = weaknesses_match.group(1).strip().replace("\n", " ")[:200] if weaknesses_match else "-"
        summary_rows.append({
            "Rank": i + 1,
            "Candidate": r["candidate_name"],
            "Final Score": round(r["final_score"], 2),
            "Fit Level": r.get("fit_level", "N/A"),
            "Top Strengths": strengths,
            "Key Weaknesses": weaknesses
        })

    summary_df = pd.DataFrame(summary_rows)
    st.dataframe(summary_df, use_container_width=True, hide_index=True)

# ------------------- EXPORT RESULTS TO WORD -------------------

if results:
    st.markdown("---")
    st.markdown("### 📥 Download All Results")

    doc = Document()

    # ----- COVER PAGE -----
    title = req_file.name if req_file else "Tender Document"
    doc.add_paragraph("")
    title_para = doc.add_paragraph("CV ASSESSMENT REPORT")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].font.size = Pt(24)
    title_para.runs[0].bold = True
    doc.add_paragraph("")
    doc.add_paragraph(f"📘 Project / Tender: {title}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"🧩 Evaluation Focus: {role_focus}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"👤 Evaluator: {evaluator_name or 'N/A'}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"📅 Date: {datetime.now().strftime('%d %B %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ----- Criteria Table -----
    if st.session_state.criteria_generated and st.session_state.custom_criteria:
        doc.add_heading("Assessment Criteria", level=2)
        table = doc.add_table(rows=1, cols=3, style="Table Grid")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Criterion'
        hdr_cells[1].text = 'Weight (%)'
        hdr_cells[2].text = 'Rationale'
        for c in st.session_state.custom_criteria["criteria"]:
            row = table.add_row().cells
            row[0].text = str(c.get("name", ""))
            row[1].text = str(c.get("weight", ""))
            row[2].text = str(c.get("rationale", ""))
        doc.add_paragraph("")

    # ----- Comparison Summary -----
    doc.add_heading("Candidate Comparison Summary", level=2)
    compare = doc.add_table(rows=1, cols=5, style="Table Grid")
    ch = compare.rows[0].cells
    ch[0].text = "Rank"; ch[1].text = "Candidate"; ch[2].text = "Final Score"
    ch[3].text = "Top Strengths"; ch[4].text = "Key Weaknesses"
    for row in summary_rows:
        r = compare.add_row().cells
        r[0].text = str(row["Rank"])
        r[1].text = row["Candidate"]
        r[2].text = str(row["Final Score"])
        r[3].text = row["Top Strengths"]
        r[4].text = row["Key Weaknesses"]
    doc.add_paragraph("")

    # ----- Detailed Reports -----
    doc.add_heading("Detailed Evaluations", level=2)

    def add_markdown_table(doc, markdown_block):
        """Convert markdown tables (|...|) into Word tables."""
        lines = [ln.strip() for ln in markdown_block.splitlines() if ln.strip()]
        if not lines or "|" not in lines[0]:
            return False
        # remove separators
        rows = [ln for ln in lines if not re.match(r"^\|[-\s|]+\|?$", ln)]
        if len(rows) < 2:
            return False
        headers = [h.strip() for h in rows[0].strip("|").split("|")]
        word_table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = word_table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                p.runs[0].bold = True
        for row in rows[1:]:
            cols = [c.strip() for c in row.strip("|").split("|")]
            new_row = word_table.add_row().cells
            for i, c in enumerate(cols):
                if i < len(new_row):
                    new_row[i].text = c
        doc.add_paragraph("")
        return True

    for i, r in enumerate(ranked):
        doc.add_heading(f"{i+1}. {r['candidate_name']}", level=3)
        clean_text = r["report"].replace("**", "").replace("#", "")
        blocks = re.split(r"(\n\|.+?\n\|[-\s|]+\|.+?(?=\n\n|\Z))", clean_text, flags=re.DOTALL)
        for block in blocks:
            if "|" in block and re.search(r"\|[-\s|]+\|", block):
                if not add_markdown_table(doc, block):
                    doc.add_paragraph(block)
            else:
                doc.add_paragraph(block.strip())
        doc.add_paragraph("-" * 80)

    # ----- SAVE FILE -----
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="📥 Download Full Assessment Report (Word)",
        data=buffer,
        file_name="CV_Assessment_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
