import os
import json
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
import PyPDF2
import mammoth
import openai

from models import CandidateAssessment


class CVAssessmentSystem:
    def __init__(self, api_key: str = None):
        """Initialize the CV assessment system"""
        from openai import OpenAI
        self.client = OpenAI(api_key=api_key or os.getenv("OPENAI_API_KEY"))
        self.job_requirements = ""
        self.assessments: List[Any] = []
        self.session_id = f"assessment_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    # ======================================================
    # 🧩 Expert Section Extraction (Hybrid: Bold + Mammoth)
    # ======================================================

    def extract_expert_sections(self, docx_path: str, target_expert_name: str) -> str:
        """
        Attempts to extract expert section text using:
        1. Bold-based detection (structured paragraphs)
        2. Fallback: Mammoth-based flat regex text
        """
        # --- 1️⃣ Try bold-based structured extraction ---
        section = self._extract_expert_sections_by_bold(docx_path, target_expert_name)
        if section and not section.startswith("⚠️") and section.strip():
            return section.strip()

        # --- 2️⃣ Fallback to mammoth-based flat text extraction ---
        flat_section = self._extract_expert_sections_by_mammoth(docx_path, target_expert_name)
        if flat_section and not flat_section.startswith("⚠️") and flat_section.strip():
            return flat_section.strip()

        return ""

    # ======================================================
    # 🧱 Primary Bold-Based Extraction
    # ======================================================

    def _extract_expert_sections_by_bold(self, docx_path: str, target_expert_name: str) -> str:
        try:
            doc = Document(docx_path)
        except Exception as e:
            return f"⚠️ Could not open document: {e}"

        # Normalize expert name and number
        target_expert_name = (
            target_expert_name.lower()
            .replace("(", "")
            .replace(")", "")
            .strip()
        )
        num_match = re.search(r"(?:key\s*expert\s*|ke\s*)(\d+)", target_expert_name, re.IGNORECASE)
        current_num = int(num_match.group(1)) if num_match else 1
        next_num = current_num + 1

        sections = []
        capture = False
        buffer = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect current expert start
            if re.search(rf"(?i)(key\s*expert\s*{current_num}\b|ke\s*{current_num}\b)", text):
                capture = True
                buffer = [text]
                continue

            # Detect next expert or section end
            if capture and re.search(
                rf"(?i)(key\s*expert\s*(?:{next_num}|[1-9]\d*)\b|ke\s*(?:{next_num}|[1-9]\d*)\b|non[-\s]*key|annex|general\s+conditions)",
                text,
            ):
                sections.append(" ".join(buffer).strip())
                buffer = []
                capture = False
                continue

            # Detect bold text marking a new header (safety cutoff)
            if capture and any(run.bold for run in para.runs if run.text.strip()):
                sections.append(" ".join(buffer).strip())
                buffer = []
                capture = False
                continue

            if capture:
                buffer.append(text)

        if buffer:
            sections.append(" ".join(buffer).strip())

        clean = [s for s in sections if len(s) > 40]
        if not clean:
            return ""
        return "\n\n---------------\n\n".join(clean)

    # ======================================================
    # 🪶 Fallback Mammoth-Based Extraction (flat text)
    # ======================================================

    def _extract_expert_sections_by_mammoth(self, docx_path: str, target_expert_name: str) -> str:
        """Fallback using flat text extraction with mammoth."""
        try:
            with open(docx_path, "rb") as f:
                result = mammoth.extract_raw_text(f)
                text = result.value
        except Exception as e:
            return f"⚠️ Could not open file: {e}"

        text = " ".join(text.split())
        if not text.strip():
            return ""

        # Normalize target expert name and number
        target_expert_name = (
            target_expert_name.lower()
            .replace("(", "")
            .replace(")", "")
            .strip()
        )
        num_match = re.search(r"(?:key\s*expert\s*|ke\s*)(\d+)", target_expert_name, re.IGNORECASE)
        current_num = int(num_match.group(1)) if num_match else 1
        next_num = current_num + 1

        # Regex to find the expert section
        pattern = re.compile(
            rf"(?i)((?:Key\s*Expert\s*{current_num}\b|KE\s*{current_num}\b).*?)"
            rf"(?=(?:Key\s*Expert\s*(?:{next_num}|[1-9]\d*)\b|KE\s*(?:{next_num}|[1-9]\d*)\b|Non[-\s]*Key|Annex|$))"
        )

        matches = pattern.findall(text)
        clean = [m.strip() for m in matches if len(m.strip()) > 30]
        if not clean:
            return ""
        return "\n\n---------------\n\n".join(clean)

    # ======================================================
    # 📄 Load Job Requirements
    # ======================================================

    def load_job_requirements(self, file_path: str) -> str:
        """Load tender or job requirements from Word or PDF file (auto-detects type)."""
        with open(file_path, "rb") as f:
            header = f.read(4)
        try:
            if header.startswith(b"%PDF"):
                text = self._extract_text_from_pdf(file_path)
            else:
                text = self._extract_text_from_word(file_path)
        except Exception as e:
            raise ValueError(f"Cannot read file {file_path}: {e}")

        self.job_requirements = text
        print(f"✅ Loaded job requirements from: {file_path}")
        return text

    # ======================================================
    # 🧠 File Helpers
    # ======================================================

    def _extract_cv_text(self, file_path: Path) -> str:
        """Extract text from CV file"""
        if file_path.suffix.lower() == ".pdf":
            return self._extract_text_from_pdf(str(file_path))
        return self._extract_text_from_word(str(file_path))

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = []
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)

    def _extract_text_from_word(self, file_path: str) -> str:
        """Extract text from Word document, preserving tables and merging broken lines."""
        doc = Document(file_path)
        lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)
        merged = []
        buffer = ""
        for line in lines:
            if len(line) < 100 and not line.endswith((".", ":", ";")):
                buffer += " " + line
            else:
                buffer += " " + line
                merged.append(buffer.strip())
                buffer = ""
        if buffer:
            merged.append(buffer.strip())
        text = "\n".join(merged)
        text = re.sub(r"\s{2,}", " ", text)
        text = re.sub(r"(\w)-\s+(\w)", r"\1\2", text)
        return text.strip()

    # ======================================================
    # 🧮 Structured Assessment Mode
    # ======================================================

    def _assess_candidate_structured(self, filename: str, cv_text: str) -> CandidateAssessment:
        """Structured scoring (dashboard mode)"""
        prompt = f"""
You are an HR evaluator performing a structured, detailed assessment of a candidate.

Compare the candidate’s CV to the job requirements below.
Assign numeric scores (0–100) and provide detailed reasoning for:
- Education
- Experience
- Skills
- Job-specific Fit

Return only valid JSON.

JOB REQUIREMENTS:
{self.job_requirements[:10000]}

CANDIDATE CV:
{cv_text[:12000]}
"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an HR evaluation system. Return ONLY valid JSON."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.1,
                max_tokens=9000,
            )

            raw_output = response.choices[0].message.content.strip()
            clean_json = self._clean_json(raw_output)
            data = json.loads(clean_json)

            return CandidateAssessment(
                candidate_name=data.get("candidate_name", filename),
                filename=filename,
                overall_score=data.get("summary", {}).get("overall_fit_score", 0),
                fit_level=data.get("summary", {}).get("fit_level", "Unknown"),
                education_details=data.get("scoring_breakdown", {}).get("education", {}),
                experience_details=data.get("scoring_breakdown", {}).get("experience", {}),
                skills_details=data.get("scoring_breakdown", {}).get("skills", {}),
                job_fit_details=data.get("scoring_breakdown", {}).get("job_specific_fit", {}),
                weighted_score_total=data.get("weighted_score_total", 0),
                executive_summary=data.get("executive_summary", {}),
                recommendation=data.get("recommendation", {}),
                interview_focus_areas=data.get("interview_focus_areas", []),
                red_flags=data.get("red_flags", []),
                potential_concerns=data.get("potential_concerns", []),
                assessed_at=datetime.now().isoformat(),
            )
        except Exception as e:
            print(f"❌ Error in structured assessment: {e}")
            return CandidateAssessment(
                candidate_name="Error",
                filename=filename,
                overall_score=0,
                fit_level="Error",
                education_details={},
                experience_details={},
                skills_details={},
                job_fit_details={},
                weighted_score_total=0,
                executive_summary={"recommendation": "Assessment failed"},
                recommendation={"verdict": "Error", "rationale": str(e)},
                interview_focus_areas=[],
                red_flags=[],
                potential_concerns=[],
                assessed_at=datetime.now().isoformat(),
            )

    # ======================================================
    # 🧭 Critical Narrative Mode
    # ======================================================

    def _assess_candidate_critical(self, filename: str, cv_text: str) -> Dict[str, Any]:
        """Critical narrative evaluation with semantic donor detection, regex fallback, and dynamic donor context."""
        donor_query = f"""
        Identify the main funding organization mentioned or implied in this tender.
        Return ONLY the donor name (e.g., 'World Bank', 'European Union', 'ADB', 'USAID', 'UNDP', 'AfDB', 'Unknown').
        If uncertain, answer exactly 'Unknown'.
        TENDER TEXT (excerpt):
        {self.job_requirements[:8000]}
        """
        donor_match = "Unknown"
        try:
            resp = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": donor_query}],
                temperature=0,
                max_tokens=10,
            )
            donor_match = resp.choices[0].message.content.strip()
        except Exception:
            pass

        donors = {
            "World Bank": r"\b(world\s*bank|wbg|ifc|ida|ibrd)\b",
            "European Union": r"\b(european\s*union|eu\s+delegation|europeaid|neighbourhood|dg\s*intl)\b",
            "Asian Development Bank": r"\b(asian\s+development\s+bank|adb)\b",
            "USAID": r"\b(usaid|united\s+states\s+agency\s+for\s+international\s+development)\b",
            "African Development Bank": r"\b(african\s+development\s+bank|afdb)\b",
            "UNDP": r"\b(undp|united\s+nations\s+development\s+programme)\b",
        }
        if donor_match == "Unknown":
            for name, pattern in donors.items():
                if re.search(pattern, self.job_requirements.lower()):
                    donor_match = name
                    break
        if donor_match == "Unknown":
            donor_match = "General donor context"

        prompt = f"""
You are a senior evaluator assessing candidates for a tender funded by **{donor_match}**.

Perform a detailed, evidence-based critical evaluation of the candidate’s CV against the JOB REQUIREMENTS and contextualize every criterion according to {donor_match}'s focus and terminology.

Return the full markdown report including:
- Evaluation Table
- Final Weighted Score
- Critical Summary
- Strengths & Weaknesses
- Tailoring Suggestions
"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a senior evaluator for international tenders. "
                            "Produce a full markdown report with scores, commentary, and recommendations."
                        ),
                    },
                    {"role": "user", "content": prompt + f"\n\nTENDER CONTEXT:\n{self.job_requirements[:4000]}\n\nCV:\n{cv_text[:9000]}"},
                ],
                temperature=0.15,
                max_tokens=8500,
            )
            report_text = response.choices[0].message.content.strip()
            match = re.search(r"Final Weighted Score.*?([0-9]\\.[0-9]+)", report_text)
            final_score = float(match.group(1)) if match else 0.0
            return {
                "candidate_name": filename,
                "report": f"**Detected Donor Context:** {donor_match}\\n\\n" + report_text,
                "final_score": final_score,
            }
        except Exception as e:
            return {"candidate_name": filename, "report": f"❌ Error generating critical evaluation: {e}", "final_score": 0.0}

    # ======================================================
    # 🧹 JSON Cleaner
    # ======================================================

    def _clean_json(self, content: str) -> str:
        """Extract clean JSON from model output."""
        content = content.strip()
        content = re.sub(r"^```(json)?", "", content)
        content = re.sub(r"```$", "", content)
        match = re.search(r"(\{.*\})", content, re.DOTALL)
        if match:
            return match.group(1).strip()
        return content
